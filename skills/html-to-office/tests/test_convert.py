import base64
import contextlib
import importlib.util
import io
import json
import os
from pathlib import Path
import shutil
import subprocess
import sys
import tempfile
import unittest

ROOT = Path(__file__).resolve().parents[1]
spec = importlib.util.spec_from_file_location("html_office_convert", ROOT / "scripts/convert.py")
convert = importlib.util.module_from_spec(spec)
spec.loader.exec_module(convert)
try:
    from lxml import html
    from PIL import Image
    PREFLIGHT = True
except ImportError:
    PREFLIGHT = False
try:
    import docx
    from pypdf import PdfReader
    OFFICE = PREFLIGHT and bool(shutil.which("pandoc"))
except ImportError:
    OFFICE = False


@unittest.skipUnless(PREFLIGHT, "requires lxml and Pillow")
class Preflight(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory(prefix="HTML 测试 ")
        self.root = Path(self.tmp.name)

    def tearDown(self):
        self.tmp.cleanup()

    def source(self, text, name="测试.html", encoding="utf-8"):
        path = self.root / name
        path.write_text(text, encoding=encoding)
        return path

    def check(self, text, **kwargs):
        return convert.prepare(self.source(text), **kwargs)

    def test_basic(self):
        report, safe, _ = self.check('<h1>中文</h1><p>正文 <b>重点</b></p>')
        self.assertFalse(report['errors'])
        self.assertIn('中文', safe)
        self.assertFalse(report['network_requested'])

    def test_dynamic_blocked(self):
        for tag in ('script', 'canvas', 'iframe', 'svg', 'math', 'video', 'form', 'custom-widget'):
            with self.subTest(tag=tag):
                self.assertTrue(self.check(f'<p>正文</p><{tag}>内容</{tag}>')[0]['errors'])

    def test_css_needs_review(self):
        source = '<style>.x{display:none}</style><p>内容</p>'
        self.assertTrue(self.check(source)[0]['errors'])
        report, _, _ = self.check(source, allow_layout_loss=True)
        self.assertFalse(report['errors'])
        self.assertTrue(report['warnings'])

    def test_layout_override_does_not_bypass_resources(self):
        report, _, _ = self.check('<style>p{color:red}</style><p>文</p><img src="missing.png">', allow_layout_loss=True)
        self.assertTrue(report['errors'])

    def test_hidden_and_tail(self):
        report, safe, _ = self.check('<p>开头<span hidden>秘密</span>结尾</p><div style="display:none">隐藏</div>')
        self.assertFalse(report['errors'])
        self.assertNotIn('秘密', safe)
        self.assertNotIn('隐藏', safe)
        self.assertIn('开头结尾', safe)

    def test_content_selection(self):
        report, safe, _ = self.check('<div>导航</div><main id="article"><p>正文</p></main>', content_id='article')
        self.assertFalse(report['errors'])
        self.assertNotIn('导航', safe)
        with self.assertRaises(convert.ConversionError):
            self.check('<p>正文</p>', content_id='missing')

    def test_dangerous_link_is_plain(self):
        report, safe, _ = self.check('<p><a href="javascript:alert(1)" onclick="alert(2)">链接</a></p>')
        self.assertNotIn('javascript:', safe)
        self.assertNotIn('onclick', safe)
        self.assertFalse(report['errors'])
        self.assertEqual(report['counts']['links'], 0)

    def test_relative_links(self):
        self.assertTrue(self.check('<a href="/docs">文档</a>')[0]['errors'])
        report, safe, _ = self.check('<a href="/docs">文档</a>', base_url='https://example.org/article')
        self.assertFalse(report['errors'])
        self.assertIn('https://example.org/docs', safe)

    def test_broken_anchor(self):
        self.assertTrue(self.check('<a href="#missing">跳转</a>')[0]['errors'])
        self.assertFalse(self.check('<h1 id="a">标题</h1><a href="#a">跳转</a>')[0]['errors'])

    def test_path_confinement(self):
        for source in ('../secret.png', '%2e%2e/secret.png', '/etc/passwd', 'file:///tmp/a.png', 'https://example.org/a.png', '//example.org/a.png'):
            with self.subTest(source=source):
                report, _, _ = self.check('<p>正文</p><img src="' + source + '">')
                self.assertTrue(report['errors'])

    def test_symlink_escape(self):
        with tempfile.TemporaryDirectory() as outside:
            image = Path(outside) / 'image.png'
            Image.new('RGB', (10, 10), 'red').save(image)
            (self.root / 'image.png').symlink_to(image)
            self.assertTrue(self.check('<img src="image.png">')[0]['errors'])

    def test_local_and_data_image(self):
        image = self.root / '图.png'
        Image.new('RGB', (40, 20), '#2070c0').save(image)
        report, safe, _ = self.check('<img src="%E5%9B%BE.png">')
        self.assertFalse(report['errors'])
        self.assertIn('data:image/png;base64,', safe)
        data = base64.b64encode(image.read_bytes()).decode('ascii')
        self.assertFalse(self.check('<img src="data:image/png;base64,' + data + '">')[0]['errors'])

    def test_animation_rejected(self):
        first = Image.new('RGB', (10, 10), 'red')
        first.save(self.root / 'animated.gif', save_all=True, append_images=[Image.new('RGB', (10, 10), 'blue')], duration=100)
        self.assertTrue(self.check('<img src="animated.gif">')[0]['errors'])

    def test_encoding(self):
        path = self.source('<p>你好</p>', encoding='gb18030')
        with self.assertRaises(convert.ConversionError):
            convert.prepare(path)
        self.assertFalse(convert.prepare(path, encoding='gb18030')[0]['errors'])

    def test_math_rejected(self):
        self.assertTrue(self.check(r'<p>\(x^2\)</p>')[0]['errors'])
        self.assertTrue(self.check('<p>$x^2$</p>')[0]['errors'])
        self.assertFalse(self.check('<p>价格 $10</p>')[0]['errors'])

    def test_pagebreak_marker(self):
        report, safe, markers = self.check('<p>首页</p><h2 style="page-break-before:always">第二页</h2>')
        self.assertFalse(report['errors'])
        self.assertEqual(len(markers), 1)
        self.assertIn(markers[0], safe)

    def test_inspect_does_not_write(self):
        source = self.source('<p>正文</p>')
        with contextlib.redirect_stdout(io.StringIO()):
            self.assertEqual(convert.main(['inspect', str(source)]), 0)
        self.assertEqual(list(self.root.iterdir()), [source])

    def test_selected_style_rejected(self):
        with self.assertRaises(convert.ConversionError):
            self.check('<style id="x">p{color:red}</style><p>内容</p>', content_id='x', allow_layout_loss=True)

    def test_complex_layout(self):
        source = '<div style="display:flex"><p>甲</p><p>乙</p></div>'
        self.assertTrue(self.check(source)[0]['errors'])
        self.assertFalse(self.check(source, allow_layout_loss=True)[0]['errors'])
        self.assertTrue(self.check('<p style="background-image:url(https://example.org/x.png)">文字</p>', allow_layout_loss=True)[0]['errors'])


@unittest.skipUnless(OFFICE, 'requires lxml, Pillow, python-docx, pypdf, Pandoc 3')
class Integration(unittest.TestCase):
    source = Preflight.source
    tearDown = Preflight.tearDown

    def setUp(self):
        Preflight.setUp(self)
        self.source_file = self.source('<h1 id="top">项目说明</h1><p>正文可编辑。</p><table><thead><tr><th>阶段</th><th>结果</th></tr></thead><tbody><tr><td rowspan="2">整理</td><td>甲</td></tr><tr><td>乙</td></tr><tr><td colspan="2">完成</td></tr></tbody></table><p><a href="https://example.org">外链</a></p><p style="page-break-before:always">第二页</p><p><img src="image.png"></p><p><a href="#top">回到开头</a></p>')
        Image.new('RGB', (240, 100), '#2674cc').save(self.root / 'image.png')

    def invoke(self, *extra):
        output = self.root / 'output'
        with contextlib.redirect_stdout(io.StringIO()):
            code = convert.main(['convert', str(self.source_file), '--output-dir', str(output), '--font', 'Arial Unicode MS', *extra])
        reports = json.loads((output / 'BATCH.json').read_text())
        self.assertEqual(code, 0, reports)
        return output, reports[0]

    def test_docx_native_structures(self):
        output, report = self.invoke('--format', 'docx')
        document = docx.Document(next(output.rglob('converted.docx')))
        self.assertIn('正文可编辑。', '\n'.join(p.text for p in document.paragraphs))
        self.assertEqual(len(document.tables), 1)
        self.assertEqual(len(document.inline_shapes), 1)
        self.assertEqual(document.tables[0].cell(1, 0).text, '整理')
        self.assertEqual(document.tables[0].cell(2, 0).text, '整理')
        self.assertEqual(document.tables[0].cell(3, 1).text, '完成')
        self.assertEqual(report['output_counts']['links'], 2)
        self.assertFalse(report['visual_reviewed'])
        self.assertFalse(list(output.rglob('*.pdf')))

    @unittest.skipUnless(shutil.which('soffice'), 'requires LibreOffice')
    def test_pdf_and_pagebreak(self):
        output, report = self.invoke('--format', 'pdf')
        reader = PdfReader(next(output.rglob('converted.pdf')))
        self.assertEqual(len(reader.pages), 2)
        self.assertIn('项目说明', reader.pages[0].extract_text())
        self.assertIn('第二页', reader.pages[1].extract_text())
        self.assertTrue(reader.pages[0].get('/Annots'))
        self.assertFalse(list(output.rglob('*.docx')))

    def test_existing_output_refused(self):
        output = self.root / 'exists'
        output.mkdir()
        (output / 'keep.txt').write_text('keep')
        with contextlib.redirect_stderr(io.StringIO()):
            code = convert.main(['convert', str(self.source_file), '--output-dir', str(output)])
        self.assertEqual(code, 2)
        self.assertEqual((output / 'keep.txt').read_text(), 'keep')

    def test_batch_partial_failure(self):
        bad = self.source('<p>缺图</p><img src="no.png">', 'bad.html')
        output = self.root / 'batch'
        with contextlib.redirect_stdout(io.StringIO()):
            code = convert.main(['convert', str(self.source_file), str(bad), '--output-dir', str(output), '--format', 'docx'])
        self.assertEqual(code, 1)
        reports = json.loads((output / 'BATCH.json').read_text())
        self.assertFalse(reports[0]['errors'])
        self.assertTrue(reports[1]['errors'])
        self.assertEqual(len(list(output.rglob('converted.docx'))), 1)

    def test_duplicate_stems(self):
        other = self.root / 'other'
        other.mkdir()
        second = other / self.source_file.name
        second.write_text('<p>第二份文件</p>', encoding='utf-8')
        output = self.root / 'batch'
        with contextlib.redirect_stdout(io.StringIO()):
            code = convert.main(['convert', str(self.source_file), str(second), '--output-dir', str(output), '--format', 'docx'])
        self.assertEqual(code, 0)
        self.assertEqual(len(list(output.rglob('converted.docx'))), 2)


class Startup(unittest.TestCase):
    def test_help_from_unrelated_directory(self):
        with tempfile.TemporaryDirectory() as tmp:
            result = subprocess.run([sys.executable, str(ROOT / 'scripts/convert.py'), '--help'], cwd=tmp, capture_output=True)
            self.assertEqual(result.returncode, 0)
            self.assertEqual(list(Path(tmp).iterdir()), [])


if __name__ == '__main__':
    unittest.main()
