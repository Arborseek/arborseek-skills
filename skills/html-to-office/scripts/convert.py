"""Offline, content-first HTML -> DOCX -> PDF. Never fetch remote resources."""
import argparse
import base64
import hashlib
import io
import json
import re
import shutil
import subprocess
import sys
import tempfile
import uuid
from pathlib import Path
from urllib.parse import unquote, urljoin, urlsplit

HTML_LIMIT = 20 * 1024 * 1024
IMAGE_LIMIT = 15 * 1024 * 1024
TOTAL_LIMIT = 80 * 1024 * 1024
ALLOWED = set("html body main article section header footer nav aside div span p br hr h1 h2 h3 h4 h5 h6 strong b em i u s del strike sub sup a img figure figcaption ul ol li dl dt dd blockquote pre code kbd samp var table caption colgroup col thead tbody tfoot tr th td abbr time mark small address center font details summary".split())
BLOCKED = set("script canvas iframe object embed video audio form input select textarea button svg math picture source".split())


class ConversionError(ValueError):
    pass


def digest(data):
    return hashlib.sha256(data).hexdigest()


def check_url(value):
    parsed = urlsplit(value)
    if parsed.scheme in ("http", "https") and parsed.hostname and not parsed.username and not parsed.password:
        return True
    return False


def raster(source, root):
    from PIL import Image
    if source.startswith("data:"):
        match = re.fullmatch(r"data:image/(?:png|jpeg|gif|webp);base64,([A-Za-z0-9+/=\s]+)", source)
        if not match:
            raise ConversionError("图片 data URI 必须是 base64 PNG/JPEG/GIF/WebP")
        try:
            data = base64.b64decode(re.sub(r"\s", "", match[1]), validate=True)
        except ValueError as exc:
            raise ConversionError("图片 base64 无效") from exc
    else:
        parsed = urlsplit(source)
        if parsed.scheme or parsed.netloc or parsed.query or parsed.fragment or not parsed.path:
            raise ConversionError("图片需本地相对路径或静态 data URI，不读取网络或 file URL")
        relative = Path(unquote(parsed.path))
        if relative.is_absolute() or "\\" in str(relative):
            raise ConversionError("图片路径不能是绝对路径或反斜杠路径")
        path = (root / relative).resolve()
        if root not in path.parents or not path.is_file():
            raise ConversionError("图片缺失或路径超出 HTML 所在目录")
        if path.stat().st_size > IMAGE_LIMIT:
            raise ConversionError("单张图片超过 15 MiB")
        data = path.read_bytes()
    if len(data) > IMAGE_LIMIT:
        raise ConversionError("单张图片超过 15 MiB")
    try:
        with Image.open(io.BytesIO(data)) as pic:
            if pic.format not in ("PNG", "JPEG", "GIF", "WEBP") or getattr(pic, "n_frames", 1) != 1:
                raise ConversionError("仅支持静态 PNG/JPEG/GIF/WebP；动态图须选定帧")
            if pic.width * pic.height > 25_000_000:
                raise ConversionError("单张图片超过 2500 万像素")
            from PIL import ImageOps
            pic = ImageOps.exif_transpose(pic).convert("RGBA")
            out = io.BytesIO()
            pic.save(out, format="PNG")
            encoded = out.getvalue()
            if len(encoded) > IMAGE_LIMIT:
                raise ConversionError("规范化后的图片超过 15 MiB")
    except ConversionError:
        raise
    except Exception as exc:
        raise ConversionError("图片损坏或无法解码") from exc
    return encoded, len(data), digest(data)


def prepare(path, *, encoding="utf-8-sig", content_id=None, base_url=None, allow_layout_loss=False):
    from lxml import etree, html
    path = Path(path).resolve()
    if path.suffix.lower() not in (".html", ".htm") or not path.is_file():
        raise ConversionError("输入必须是存在的本地 .html/.htm 文件")
    if path.stat().st_size > HTML_LIMIT:
        raise ConversionError("HTML 超过 20 MiB")
    raw = path.read_bytes()
    try:
        source = raw.decode(encoding, errors="strict")
    except (UnicodeError, LookupError) as exc:
        raise ConversionError("无法按指定编码解码，请检查 --encoding") from exc
    if base_url and not check_url(base_url):
        raise ConversionError("--base-url 必须是无账号密码的 HTTP(S) URL")
    try:
        html_parser = html.HTMLParser(no_network=True, remove_comments=True)
        tree = html.document_fromstring(source, parser=html_parser)
        if any(item.level_name == "FATAL" for item in html_parser.error_log):
            raise ConversionError("HTML 解析发生严重错误，可能截断正文")
    except (etree.ParserError, ValueError) as exc:
        raise ConversionError("无法解析 HTML") from exc
    if content_id:
        matches = tree.xpath("//*[@id=$target]", target=content_id)
        if len(matches) != 1:
            raise ConversionError("--content-id 必须匹配唯一容器")
        body = matches[0]
    else:
        body = tree.find("body")
    if body is None:
        raise ConversionError("HTML 没有正文")
    if body.tag not in ALLOWED:
        raise ConversionError("选定容器不是可转换的正文元素")
    errors, warnings = [], []
    css = tree.xpath("//style | //link[contains(translate(@rel,'ABCDEFGHIJKLMNOPQRSTUVWXYZ','abcdefghijklmnopqrstuvwxyz'),'stylesheet')]")
    if css:
        if not allow_layout_loss:
            errors.append("含样式表：检查 CSS 隐藏内容及阅读顺序后才可使用 --allow-layout-loss")
        else:
            warnings.append("已按调用者确认忽略样式表；PDF/Word 使用文档布局")
    if tree.xpath("//base"):
        warnings.append("忽略 HTML base；相对链接仅按显式 --base-url 解析")
    # Scripts anywhere can populate the selected content; never pretend to render them.
    if tree.xpath("//script"):
        errors.append("含 script 动态内容或脚本：请先导出静态正文")
    for node in list(body.iterdescendants()):
        if not isinstance(node.tag, str) or node.getparent() is None:
            continue
        style = node.get("style", "")
        if "hidden" in node.attrib or re.search(r"(?:^|;)\s*(?:display\s*:\s*none|visibility\s*:\s*hidden)\s*(?:!important)?\s*(?:;|$)", style, re.I):
            node.drop_tree()
    if "hidden" in body.attrib or re.search(r"display\s*:\s*none|visibility\s*:\s*hidden", body.get("style", ""), re.I):
        raise ConversionError("选定正文容器不可见")
    images, links, page_markers, total = [], 0, [], 0
    for node in list(body.iter()):
        if not isinstance(node.tag, str):
            continue
        tag = node.tag.lower()
        if tag in ("style", "link"):
            if node is not body:
                node.drop_tree()
            continue
        if tag in BLOCKED or tag not in ALLOWED:
            errors.append("不支持的内容 <" + tag + ">：需要静态替代或明确处理")
            continue
        if tag in ("details", "summary"):
            errors.append("含折叠内容 details/summary：请先确定需要展开的正文")
        style = node.get("style", "")
        if style:
            warnings.append("内联字体、颜色和间距规范化为文档样式")
        if re.search(r"url\s*\(|content\s*:", style, re.I):
            errors.append("含 CSS 图片/生成内容：需先转为真实正文或 img")
        if re.search(r"(?:display\s*:\s*(?:grid|flex)|position\s*:\s*(?:absolute|fixed)|(?:column-count|order)\s*:)", style, re.I):
            if not allow_layout_loss:
                errors.append("含复杂内联布局：检查阅读顺序后才可使用 --allow-layout-loss")
        for rule, position in ((r"(?:page-break-before|break-before)\s*:\s*(?:always|page)", "before"), (r"(?:page-break-after|break-after)\s*:\s*(?:always|page)", "after")):
            if re.search(rule, style, re.I):
                if tag not in ("p", "div", "section", "h1", "h2", "h3", "h4", "h5", "h6", "table", "article") or node is body or node.xpath("ancestor::table"):
                    errors.append("此元素上的分页无法可靠转换，请将分页放在顶层段落")
                else:
                    token = "OFFICEPAGEBREAK" + uuid.uuid4().hex
                    marker = html.Element("p")
                    marker.text = token
                    if position == "before":
                        node.addprevious(marker)
                    else:
                        node.addnext(marker)
                    page_markers.append(token)
        attrs = {}
        if node.get("id"):
            attrs["id"] = node.get("id")
        if tag == "a" and node.get("href"):
            href = node.get("href").strip()
            if href.startswith("#") or check_url(href) or (href.startswith("mailto:") and not re.search(r"[\r\n]", href)):
                attrs["href"] = href
            elif not urlsplit(href).scheme:
                if base_url and check_url(urljoin(base_url, href)):
                    attrs["href"] = urljoin(base_url, href)
                else:
                    errors.append("相对超链接缺少有效 --base-url")
            else:
                warnings.append("危险或不支持的超链接已降级为普通文字")
            if "href" in attrs:
                links += 1
        if tag == "img":
            if node.get("srcset"):
                errors.append("含响应式 srcset：请先选定一张图片")
            try:
                data, size, sha = raster(node.get("src", ""), path.parent)
                total += max(size, len(data))
                if total > TOTAL_LIMIT:
                    raise ConversionError("图片总量超过 80 MiB")
                attrs["src"] = "data:image/png;base64," + base64.b64encode(data).decode("ascii")
                attrs["alt"] = node.get("alt", "")
                images.append(sha)
            except ConversionError as exc:
                errors.append("图片 " + str(len(images) + 1) + "：" + str(exc))
        for key in (("colspan", "rowspan") if tag in ("td", "th") else ("start",) if tag == "ol" else ()):
            if node.get(key):
                if not node.get(key).isdigit() or not 1 <= int(node.get(key)) <= 1000:
                    errors.append("表格跨度/列表起点必须为 1–1000 的整数")
                else:
                    attrs[key] = node.get(key)
        node.attrib.clear()
        node.attrib.update(attrs)
    text = body.text_content().strip()
    if re.search(r"\\\(|\\\[|\$\$|(?<![\\\w])\$(?!\d)[^$\n]+\$", text):
        errors.append("含未渲染 LaTeX 公式：请先转换为可验证的公式或图片")
    if not text and not images:
        errors.append("没有可转换的正文或图片")
    ids = [node.get("id") for node in body.iter() if node.get("id")]
    if len(set(ids)) != len(ids):
        errors.append("正文包含重复 ID，内部跳转不可靠")
    for node in body.xpath(".//a[starts-with(@href,'#')]"):
        if unquote(node.get("href")[1:]) not in ids:
            errors.append("内部链接没有对应锚点")
    result = {"input": str(path), "source_sha256": digest(raw), "warnings": sorted(set(warnings)),
              "errors": sorted(set(errors)), "counts": {"images": len(images), "tables": len(body.xpath(".//table")) + (body.tag == "table"), "links": links},
              "image_sha256": images, "network_requested": False, "visual_reviewed": False}
    # Only sanitized static content, never head/base or hidden source instructions.
    safe = "<!doctype html><html><head><meta charset='utf-8'></head><body>" + html.tostring(body, encoding="unicode") + "</body></html>"
    return result, safe, page_markers


def executable(value, fallback):
    path = shutil.which(value or fallback)
    if not path:
        raise ConversionError("缺少转换工具：" + fallback)
    return path


def run(command, cwd):
    try:
        proc = subprocess.run(command, cwd=cwd, capture_output=True, text=True, encoding="utf-8", errors="replace", timeout=120)
    except (OSError, subprocess.TimeoutExpired) as exc:
        raise ConversionError("转换工具无法启动或超过 120 秒") from exc
    if proc.returncode:
        raise ConversionError("转换工具失败（退出码 " + str(proc.returncode) + "），请检查工具版本与环境")
    return proc


def style_docx(path, font, paper, landscape, markers):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.enum.text import WD_BREAK
    from docx.enum.section import WD_ORIENT
    document = Document(path)
    width, height = (8.27, 11.69) if paper == "a4" else (8.5, 11)
    if landscape:
        width, height = height, width
    for section in document.sections:
        section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
        section.page_width, section.page_height = Inches(width), Inches(height)
        section.top_margin = section.bottom_margin = Inches(.7)
        section.left_margin = section.right_margin = Inches(.75)
    for style in document.styles:
        if hasattr(style, "font"):
            style.font.name = font
            style.font.size = Pt(11)
            style.font.color.rgb = RGBColor(32, 38, 46)
            rpr = style.element.get_or_add_rPr()
            fonts = rpr.find(qn("w:rFonts"))
            if fonts is None:
                fonts = OxmlElement("w:rFonts")
                rpr.append(fonts)
            for key in list(fonts.attrib):
                if key.endswith("Theme"):
                    del fonts.attrib[key]
            fonts.set(qn("w:eastAsia"), font)
        ppr = style.element.find(qn("w:pPr"))
        if ppr is not None:
            for border in list(ppr.findall(qn("w:pBdr"))):
                ppr.remove(border)
    for name, size in (("Title", 22), ("Heading 1", 20), ("Heading 2", 16), ("Heading 3", 13)):
        if name in document.styles:
            style = document.styles[name]
            style.font.size, style.font.bold = Pt(size), True
            style.font.color.rgb = RGBColor(0, 0, 0)
            style.paragraph_format.space_before = Pt(12)
            style.paragraph_format.space_after = Pt(7)
    for name in ("Normal", "Body Text", "First Paragraph", "Compact"):
        if name in document.styles:
            document.styles[name].paragraph_format.line_spacing = 1.2
            document.styles[name].paragraph_format.space_after = Pt(7)
    if "Hyperlink" in document.styles:
        document.styles["Hyperlink"].font.color.rgb = RGBColor(30, 93, 166)
        document.styles["Hyperlink"].font.underline = True
    for p in document.paragraphs:
        if p.text.strip() in markers:
            for child in list(p._p):
                p._p.remove(child)
            p.add_run().add_break(WD_BREAK.PAGE)
    usable = Inches(width - 1.5)
    for table in document.tables:
        table.autofit = True
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = OxmlElement("w:" + edge)
            for key, value in (("val", "single"), ("sz", "4"), ("color", "CDD5DF")):
                border.set(qn("w:" + key), value)
            borders.append(border)
        table._tbl.tblPr.append(borders)
        grid = table._tbl.tblGrid
        old = [int(col.get(qn("w:w"))) for col in grid]
        ratio = usable.twips / max(sum(old), 1)
        for col, value in zip(grid, old):
            col.set(qn("w:w"), str(max(1, round(value * ratio))))
        for cell in table._tbl.iter(qn("w:tc")):
            tcpr = cell.find(qn("w:tcPr"))
            tcw = tcpr.find(qn("w:tcW")) if tcpr is not None else None
            if tcw is not None and tcw.get(qn("w:type")) == "dxa":
                tcw.set(qn("w:w"), str(max(1, round(int(tcw.get(qn("w:w"))) * ratio))))
        if table.rows:
            first = table.rows[0]._tr
            # Repeat only an actual header recognized by Pandoc, don't invent headers.
            header = first.find(".//" + qn("w:tblHeader"))
            if header is not None:
                for cell in table.rows[0].cells:
                    shade = OxmlElement("w:shd")
                    shade.set(qn("w:fill"), "EDF2F7")
                    cell._tc.get_or_add_tcPr().append(shade)
    for shape in document.inline_shapes:
        factor = min(1, usable / shape.width, Inches(height - 1.8) / shape.height)
        shape.width, shape.height = round(shape.width * factor), round(shape.height * factor)
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    document.save(path)
    # No linked images, templates, OLE, or active content may enter the office package.
    import zipfile
    from lxml import etree
    counts = {"images": 0, "tables": 0, "links": 0}
    with zipfile.ZipFile(path) as archive:
        for name in archive.namelist():
            if "embeddings/" in name or "vbaProject" in name:
                raise ConversionError("输出包含不允许的嵌入对象")
            if name.endswith(".rels"):
                for rel in etree.fromstring(archive.read(name)):
                    if rel.get("TargetMode") == "External" and not rel.get("Type", "").endswith("/hyperlink"):
                        raise ConversionError("输出包含非超链接外部关系")
        xml = etree.fromstring(archive.read("word/document.xml"))
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        for key, tag in (("images", "drawing"), ("tables", "tbl"), ("links", "hyperlink")):
            counts[key] = len(xml.xpath("//w:" + tag, namespaces=ns))
    return counts


def convert_one(source, target, args):
    report, safe, markers = prepare(source, encoding=args.encoding, content_id=args.content_id,
                                    base_url=args.base_url, allow_layout_loss=args.allow_layout_loss)
    if report["errors"]:
        return report
    pandoc = executable(args.pandoc, "pandoc")
    version = run([pandoc, "--version"], target).stdout.splitlines()[0]
    found_version = re.search(r"pandoc (\d+)\.", version)
    if not found_version or int(found_version[1]) < 3:
        raise ConversionError("需要 Pandoc 3.x 或更新版本（支持 --sandbox）")
    public = target / "public"
    with tempfile.TemporaryDirectory(prefix="html-office-") as tmp:
        work = Path(tmp)
        html_path = work / "sanitized.html"
        html_path.write_text(safe, encoding="utf-8")
        docx = work / "converted.docx"
        data_dir = work / "pandoc-data"
        data_dir.mkdir()
        proc = run([pandoc, "--sandbox", "--data-dir", str(data_dir), "--from=html", "--to=docx", str(html_path), "--output", str(docx)], work)
        if proc.stderr.strip():
            raise ConversionError("Pandoc 报告警告；本次不交付可能缺失内容的文件，请检查输入支持范围")
        counts = style_docx(docx, args.font, args.paper, args.landscape, markers)
        if counts != report["counts"]:
            raise ConversionError("转换前后图片/表格/链接数量不一致，停止交付")
        report["output_counts"] = counts
        report["backend"] = {"docx": version, "pdf": "LibreOffice Writer" if args.format != "docx" else None}
        deliver = [docx] if args.format in ("docx", "both") else []
        if args.format in ("pdf", "both"):
            soffice = executable(args.soffice, "soffice")
            pdf_dir = work / "pdf"
            pdf_dir.mkdir()
            profile = work / "lo-profile"
            run([soffice, "-env:UserInstallation=" + profile.as_uri(), "--headless", "--convert-to", "pdf:writer_pdf_Export", "--outdir", str(pdf_dir), str(docx)], work)
            pdf = pdf_dir / "converted.pdf"
            if not pdf.is_file() or not pdf.read_bytes().startswith(b"%PDF-"):
                raise ConversionError("PDF 引擎未生成有效文件")
            from pypdf import PdfReader
            reader = PdfReader(pdf)
            if not reader.pages:
                raise ConversionError("PDF 没有页面")
            report["pdf_pages"] = len(reader.pages)
            report["pdf_text_characters"] = sum(len(p.extract_text() or "") for p in reader.pages)
            deliver.append(pdf)
        public.mkdir()
        report["outputs"] = []
        for path in deliver:
            destination = public / path.name
            shutil.copyfile(path, destination)
            report["outputs"].append({"path": str(destination), "sha256": digest(destination.read_bytes())})
    report["status"] = "converted_pending_visual_review"
    return report


def parser():
    parser = argparse.ArgumentParser(description=__doc__)
    sub = parser.add_subparsers(dest="command", required=True)
    for command in ("inspect", "convert"):
        part = sub.add_parser(command)
        part.add_argument("inputs", nargs="+", type=Path)
        part.add_argument("--encoding", default="utf-8-sig")
        part.add_argument("--content-id")
        part.add_argument("--base-url")
        part.add_argument("--allow-layout-loss", action="store_true")
        if command == "convert":
            part.add_argument("--output-dir", type=Path, required=True)
            part.add_argument("--format", choices=("docx", "pdf", "both"), default="both")
            part.add_argument("--font", default="Noto Sans CJK SC")
            part.add_argument("--paper", choices=("letter", "a4"), default="letter")
            part.add_argument("--landscape", action="store_true")
            part.add_argument("--pandoc")
            part.add_argument("--soffice")
    return parser


def main(argv=None):
    args = parser().parse_args(argv)
    if args.command == "convert":
        output = args.output_dir.absolute()
        try:
            output.mkdir(parents=True, exist_ok=False)
        except FileExistsError:
            print("输出目录已存在；请使用新目录", file=sys.stderr)
            return 2
    reports = []
    for index, source in enumerate(args.inputs, 1):
        target = None
        try:
            if args.command == "inspect":
                report, _, _ = prepare(source, encoding=args.encoding, content_id=args.content_id, base_url=args.base_url, allow_layout_loss=args.allow_layout_loss)
                report["status"] = "blocked" if report["errors"] else "ready"
            else:
                stem = re.sub(r"[^\w-]", "-", source.stem)[:60] or "document"
                target = output / (f"{index:03d}-" + stem)
                target.mkdir()
                report = convert_one(source, target, args)
                if report["errors"]:
                    report["status"] = "blocked"
        except (ConversionError, OSError, ImportError) as exc:
            report = {"input": str(source), "status": "failed", "errors": [str(exc)], "visual_reviewed": False}
        except Exception as exc:
            report = {"input": str(source), "status": "failed", "errors": ["转换未完成：" + type(exc).__name__ + "；请检查依赖与输入支持范围"], "visual_reviewed": False}
        reports.append(report)
        if target:
            (target / "report.json").write_text(json.dumps(report, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    if args.command == "convert":
        (output / "BATCH.json").write_text(json.dumps(reports, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(json.dumps(reports, ensure_ascii=False, indent=2))
    return int(any(r.get("errors") for r in reports))


if __name__ == "__main__":
    sys.exit(main())
