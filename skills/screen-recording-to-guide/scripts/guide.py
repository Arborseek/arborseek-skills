"""Check reviewed evidence and export a screenshot guide. No AI or network calls."""
import argparse
import copy
import html
import json
import math
import re
from pathlib import Path

from capture import digest, finite


def nonempty(value):
    return isinstance(value, str) and bool(value.strip())


def strings(value):
    return isinstance(value, list) and all(nonempty(s) for s in value)


def rect(box, width, height):
    return (isinstance(box, list) and len(box) == 4 and all(isinstance(v, int) and not isinstance(v, bool) for v in box)
            and box[0] >= 0 and box[1] >= 0 and box[2] > 0 and box[3] > 0
            and box[0] + box[2] <= width and box[1] + box[3] <= height)


def arrow(points, width, height):
    return (isinstance(points, list) and len(points) == 4 and all(isinstance(v, int) and not isinstance(v, bool) for v in points)
            and 0 <= points[0] < width and 0 <= points[2] < width
            and 0 <= points[1] < height and 0 <= points[3] < height and points[:2] != points[2:])


def inside_box(box, crop):
    x, y, w, h = crop
    return x <= box[0] and y <= box[1] and box[0] + box[2] <= x + w and box[1] + box[3] <= y + h


def load_sources(plan, base):
    from PIL import Image
    sources, errors = {}, []
    aliases = plan.get('sources')
    if not isinstance(aliases, dict) or not aliases:
        return {}, ['sources must map aliases to extraction indexes']
    hashes = set()
    for alias, value in aliases.items():
        try:
            if not nonempty(alias) or not nonempty(value):
                raise ValueError('Invalid source alias or path')
            index_path = (base / value).resolve(strict=True)
            index = json.loads(index_path.read_text(encoding='utf-8'))
            if not isinstance(index, dict) or index.get('schema') != 'screen-frames/1':
                raise ValueError('Invalid index schema')
            source = index.get('source', {})
            if not isinstance(source, dict) or not re.fullmatch(r'[a-f0-9]{64}', str(source.get('sha256', ''))):
                raise ValueError('Invalid video hash')
            duration = source.get('duration')
            if not finite(duration) or duration <= 0:
                raise ValueError('Invalid video duration')
            hashes.add(source['sha256'])
            frames = index.get('frames')
            if not isinstance(frames, list) or not frames:
                raise ValueError('Empty frames index')
            by_id = {}
            for frame in frames:
                if not isinstance(frame, dict) or not nonempty(frame.get('id')) or frame['id'] in by_id:
                    raise ValueError('Invalid or duplicate frame ID')
                if not finite(frame.get('at')) or not 0 <= frame['at'] < duration:
                    raise ValueError('Frame time out of range')
                if not nonempty(frame.get('file')):
                    raise ValueError('Missing frame path')
                path = (index_path.parent / frame['file']).resolve(strict=True)
                if index_path.parent not in path.parents or path.suffix.lower() != '.png':
                    raise ValueError('Frame must be a PNG inside its extraction directory')
                if digest(path) != frame.get('sha256'):
                    raise ValueError('Frame hash mismatch')
                with Image.open(path) as image:
                    if image.format != 'PNG' or image.size != (frame.get('width'), frame.get('height')):
                        raise ValueError('Frame dimensions or format mismatch')
                    image.verify()
                by_id[frame['id']] = {**frame, 'path': path}
            sources[alias] = {'index': index, 'frames': by_id}
        except (ValueError, OSError, KeyError, TypeError) as exc:
            errors.append(f'source {alias}: {exc}')
    if len(hashes) > 1:
        errors.append('Extraction indexes refer to different videos')
    return sources, errors


def validate(plan, base):
    errors, blockers = [], []
    if not isinstance(plan, dict):
        return {'valid': False, 'ready': False, 'errors': ['plan must be an object'], 'blockers': []}, {}
    if plan.get('schema') != 'screen-guide/1':
        errors.append('Unsupported schema')
    for key in ('title', 'purpose', 'audience', 'scope'):
        if not nonempty(plan.get(key)):
            errors.append(key + ' must be a nonempty string')
    for key in ('prerequisites', 'completion', 'questions'):
        if not strings(plan.get(key)):
            errors.append(key + ' must be a string array')
    if plan.get('coverage_reviewed') is not True:
        blockers.append('录屏流程覆盖和前后衔接尚未核对')
    if isinstance(plan.get('questions'), list):
        blockers.extend(q for q in plan['questions'] if nonempty(q))
    sources, source_errors = load_sources(plan, base)
    errors.extend(source_errors)
    steps = plan.get('steps')
    if not isinstance(steps, list) or not steps:
        errors.append('steps must be a nonempty array')
        steps = []
    ids = set()
    previous_time = -1
    for n, step in enumerate(steps, 1):
        prefix = f'step {n}: '
        if not isinstance(step, dict):
            errors.append(prefix + 'must be an object')
            continue
        for key in ('id', 'title', 'instruction', 'expected'):
            if not nonempty(step.get(key)):
                errors.append(prefix + key + ' is required')
        ident = step.get('id')
        if nonempty(ident):
            if ident in ids:
                errors.append(prefix + 'duplicate ID')
            ids.add(ident)
        if step.get('verified') is not True:
            blockers.append(f'第 {n} 步动作与结果尚未核对')
        if step.get('privacy_reviewed') is not True:
            blockers.append(f'第 {n} 步图片与文字隐私尚未核对')
        evidence = step.get('evidence')
        if not isinstance(evidence, list) or not evidence:
            errors.append(prefix + 'at least one evidence frame is required')
            continue
        step_times = []
        for item in evidence:
            if not isinstance(item, dict) or not nonempty(item.get('source')) or not nonempty(item.get('frame')):
                errors.append(prefix + 'invalid evidence reference')
                continue
            frame = sources.get(item['source'], {}).get('frames', {}).get(item['frame'])
            if frame is None:
                errors.append(prefix + 'unknown frame reference')
                continue
            step_times.append(frame['at'])
            if item.get('role') not in ('before', 'after', 'detail'):
                errors.append(prefix + 'invalid evidence role')
            w, h = frame['width'], frame['height']
            crop = item.get('crop', [0, 0, w, h])
            if not rect(crop, w, h):
                errors.append(prefix + 'invalid crop')
                continue
            for field in ('redactions', 'boxes', 'arrows'):
                shapes = item.get(field, [])
                if not isinstance(shapes, list):
                    errors.append(prefix + field + ' must be an array')
                    continue
                for shape in shapes:
                    valid = arrow(shape, w, h) if field == 'arrows' else rect(shape, w, h)
                    if not valid:
                        errors.append(prefix + 'invalid ' + field)
                    elif field == 'boxes' and not inside_box(shape, crop):
                        errors.append(prefix + 'crop would remove a box')
                    elif field == 'arrows' and not all(crop[0] <= shape[i] < crop[0] + crop[2] and crop[1] <= shape[i+1] < crop[1] + crop[3] for i in (0, 2)):
                        errors.append(prefix + 'crop would remove an arrow')
        if step_times:
            if min(step_times) < previous_time:
                blockers.append(f'第 {n} 步时间倒序，请重排步骤或核对是否引用了错误片段')
            previous_time = min(step_times)
    return {'valid': not errors, 'ready': not errors and not blockers, 'errors': errors, 'blockers': blockers}, sources


def annotate(frame, evidence, target):
    from PIL import Image, ImageDraw, ImageFont
    with Image.open(frame['path']) as original:
        # Paste onto a fresh image to drop source metadata and alpha channels.
        image = Image.new('RGB', original.size, 'white')
        image.paste(original.convert('RGB'))
    draw = ImageDraw.Draw(image)
    for x, y, w, h in evidence.get('redactions', []):
        draw.rectangle((x, y, x+w-1, y+h-1), fill='#111111')
    stroke = max(3, round(min(image.size) / 200))
    for n, (x, y, w, h) in enumerate(evidence.get('boxes', []), 1):
        draw.rectangle((x, y, x+w-1, y+h-1), outline='#dc2626', width=stroke)
        # Number above the object when possible; do not obscure the label.
        crop = evidence.get('crop', [0, 0, image.width, image.height])
        badge_y = y - 21 if y - 21 >= crop[1] else y
        draw.rectangle((x, badge_y, x+20, badge_y+19), fill='#dc2626')
        draw.text((x+5, badge_y+3), str(n), fill='white', font=ImageFont.load_default())
    for x1, y1, x2, y2 in evidence.get('arrows', []):
        draw.line((x1, y1, x2, y2), fill='#dc2626', width=stroke)
        angle = math.atan2(y2-y1, x2-x1)
        length = stroke * 4
        points = [(x2, y2)] + [(x2-length*math.cos(angle+a), y2-length*math.sin(angle+a)) for a in (-.5, .5)]
        draw.polygon(points, fill='#dc2626')
    if 'crop' in evidence:
        x, y, w, h = evidence['crop']
        image = image.crop((x, y, x+w, y+h))
    image.save(target, format='PNG')


def clock(at):
    ms = round(at * 1000)
    hours, rest = divmod(ms, 3600000)
    minutes, rest = divmod(rest, 60000)
    secs, ms = divmod(rest, 1000)
    return f'{hours:02}:{minutes:02}:{secs:02}.{ms:03}'


def safe_md(text):
    return html.escape(text).replace('\\', '\\\\').replace('[', '\\[').replace(']', '\\]').replace('*', '\\*').replace('#', '\\#')


def docx_export(plan, pages, path, draft, font):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.oxml.ns import qn
    from PIL import Image
    doc = Document()
    section = doc.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.bottom_margin = Inches(.75)
    section.left_margin = section.right_margin = Inches(.8)
    for border in list(doc.styles.element.iter(qn('w:pBdr'))):
        border.getparent().remove(border)
    for name in ('Normal', 'Title', 'Heading 1', 'Heading 2', 'List Bullet'):
        style = doc.styles[name]
        style.font.name = font
        style.font.color.rgb = RGBColor(0, 0, 0)
        fonts = style.element.get_or_add_rPr().rFonts
        for attr in list(fonts.attrib):
            if attr.endswith('Theme'):
                del fonts.attrib[attr]
        for attr in ('ascii', 'hAnsi', 'eastAsia', 'cs'):
            fonts.set(qn('w:' + attr), font)
    doc.styles['Normal'].font.size = Pt(11)
    doc.styles['Normal'].paragraph_format.space_after = Pt(8)
    doc.styles['Normal'].paragraph_format.line_spacing = 1.15
    doc.styles['Title'].font.size = Pt(24)
    doc.styles['Heading 1'].font.size = Pt(16)
    doc.add_heading(plan['title'], 0)
    if draft:
        doc.add_paragraph('内部预览，待核对后使用。')
    doc.add_paragraph(plan['purpose'])
    doc.add_paragraph('适用对象：' + plan['audience'] + '。适用范围：' + plan['scope'])
    if plan['prerequisites']:
        doc.add_heading('准备工作', 1)
        for value in plan['prerequisites']:
            doc.add_paragraph(value, 'List Bullet')
    for number, (step, pictures) in enumerate(pages, 1):
        doc.add_heading(f"{number} {step['title']}", 1)
        paragraph = doc.add_paragraph(step['instruction'])
        paragraph.paragraph_format.keep_with_next = True
        for picture in pictures:
            with Image.open(picture['path']) as image:
                w, h = image.size
            # Avoid giant portrait screenshots exceeding a page.
            width = min(6.9, 4.1 * w / h)
            p = doc.add_paragraph()
            p.paragraph_format.keep_with_next = True
            shape = p.add_run().add_picture(str(picture['path']), width=Inches(width))
            shape._inline.docPr.set('descr', step['title'] + ' ' + picture['caption'])
            caption = doc.add_paragraph(picture['caption'])
            caption.paragraph_format.space_after = Pt(7)
        doc.add_paragraph('完成检查：' + step['expected'])
    if plan['completion']:
        doc.add_heading('最终检查', 1)
        for value in plan['completion']:
            doc.add_paragraph(value, 'List Bullet')
    doc.core_properties.author = ''
    doc.core_properties.last_modified_by = ''
    doc.save(path)


def build(plan_path, output, draft=False, docx=False, font='Noto Sans CJK SC'):
    plan_path = Path(plan_path).resolve(strict=True)
    plan = json.loads(plan_path.read_text(encoding='utf-8'))
    report, sources = validate(plan, plan_path.parent)
    if not report['valid'] or (not draft and not report['ready']):
        raise ValueError(json.dumps(report, ensure_ascii=False))
    if docx:
        import docx as docx_library  # fail before creating output when unavailable
    output = Path(output)
    output.mkdir(parents=True, exist_ok=False)
    public = output / 'public'
    (public / 'images').mkdir(parents=True)
    esc = html.escape
    md = ['# ' + safe_md(plan['title']), '']
    html_parts = ['<!doctype html><html lang="zh-CN"><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">',
                  '<title>' + esc(plan['title']) + '</title>',
                  '<style>body{font:17px/1.7 system-ui,sans-serif;color:#18232d;max-width:900px;margin:40px auto;padding:0 24px}h1,h2{line-height:1.35;color:#101820}h1{font-size:30px}h2{font-size:23px;margin-top:36px}img{max-width:100%;height:auto;border:1px solid #dfe5e9}figure{margin:22px 0}figcaption{font-size:13px;color:#586574}a{color:#175e83}p{white-space:pre-wrap}.draft{color:#a43c00}@media print{nav{display:none}body{margin:0;font-size:11pt}h2{break-after:avoid}figure{break-inside:avoid}img{max-height:180mm;object-fit:contain}}</style>',
                  '<body><h1>' + esc(plan['title']) + '</h1>']
    if draft:
        md.extend(['内部预览，待核对后使用。', ''])
        html_parts.append('<p class="draft">内部预览，待核对后使用。</p>')
    for text in (plan['purpose'], '适用对象：' + plan['audience'], '适用范围：' + plan['scope']):
        md.extend([safe_md(text), ''])
        html_parts.append('<p>' + esc(text) + '</p>')
    if plan['prerequisites']:
        md.extend(['## 准备工作', ''] + ['- ' + safe_md(x) for x in plan['prerequisites']] + [''])
        html_parts.append('<h2>准备工作</h2><ul>' + ''.join('<li>' + esc(x) + '</li>' for x in plan['prerequisites']) + '</ul>')
    html_parts.append('<nav><ol>' + ''.join(f'<li><a href="#step-{i}">{esc(s["title"])}</a></li>' for i, s in enumerate(plan['steps'], 1)) + '</ol></nav>')
    pages, evidence_log = [], []
    roles = {'before': '操作前', 'after': '操作后', 'detail': '细节'}
    for number, step in enumerate(plan['steps'], 1):
        md.extend([f"## {number} {safe_md(step['title'])}", '', safe_md(step['instruction']), ''])
        html_parts.append(f'<section id="step-{number}"><h2>{number} {esc(step["title"])}</h2><p>{esc(step["instruction"])}</p>')
        pictures = []
        for count, item in enumerate(step['evidence'], 1):
            frame = sources[item['source']]['frames'][item['frame']]
            relative = f'images/step-{number:03}-{count:02}.png'
            path = public / relative
            annotate(frame, item, path)
            caption = f"{roles[item['role']]} · 原视频 {clock(frame['at'])}"
            pictures.append({'path': path, 'caption': caption})
            md.extend([f'![{safe_md(step["title"])}]({relative})', '', caption, ''])
            html_parts.append(f'<figure><img src="{relative}" alt="{esc(step["title"], quote=True)}"><figcaption>{esc(caption)}</figcaption></figure>')
            evidence_log.append({'step': step['id'], 'reference': item, 'source_frame_sha256': frame['sha256'],
                                 'requested_seconds': frame['at'], 'public_file': relative, 'output_sha256': digest(path)})
        md.extend(['完成检查：' + safe_md(step['expected']), ''])
        html_parts.append('<p>完成检查：' + esc(step['expected']) + '</p></section>')
        pages.append((step, pictures))
    if plan['completion']:
        md.extend(['## 最终检查', ''] + ['- ' + safe_md(x) for x in plan['completion']] + [''])
        html_parts.append('<h2>最终检查</h2><ul>' + ''.join('<li>' + esc(x) + '</li>' for x in plan['completion']) + '</ul>')
    html_parts.append('</body></html>')
    (public / 'tutorial.html').write_text('\n'.join(html_parts), encoding='utf-8')
    (public / 'tutorial.md').write_text('\n'.join(md), encoding='utf-8')
    if docx:
        docx_export(plan, pages, public / 'tutorial.docx', draft, font)
    notes = ['# 内部核对单', '', '这是制作记录，不随 public 目录外发。', '',
             '导出模式：' + ('内部预览' if draft else '已通过输入核对门槛，仍需交付验收'), '', '## 待确认', '']
    notes.extend(['- ' + q for q in report['blockers']] or ['无登记的未决项；这不是对语义正确性的自动认证。'])
    notes.extend(['', '## 交付验收', '', '- 逐步核对正文、截图、按钮标注及操作结果。',
                  '- 检查所有截图和文字中的账号、个人信息与业务机密。',
                  '- 检查 HTML / Word 中的字体、图片可读性及分页。',
                  '- 请不熟悉流程的员工试走；未做则不要声称已通过实际操作验收。',
                  '- 只分享 public，不携带本核对单、溯源记录或原视频。', ''])
    (output / 'review.md').write_text('\n'.join(notes), encoding='utf-8')
    provenance = {'plan': copy.deepcopy(plan), 'sources': {k: v['index'] for k, v in sources.items()}, 'evidence': evidence_log}
    (output / 'provenance.json').write_text(json.dumps(provenance, ensure_ascii=False, indent=2) + '\n', encoding='utf-8')
    files = {str(p.relative_to(output)): digest(p) for p in output.rglob('*') if p.is_file()}
    receipt = {'complete': True, 'mode': 'draft' if draft else 'reviewed-input', 'steps': len(pages), 'files': files,
               'semantic_accuracy_certified': False, 'visual_export_reviewed': False}
    (output / 'COMPLETE.json').write_text(json.dumps(receipt, ensure_ascii=False, indent=2) + '\n', encoding='utf-8')
    return {'ok': True, 'steps': len(pages), 'draft': draft, 'docx': docx}


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument('command', choices=('check', 'build'))
    parser.add_argument('plan', type=Path)
    parser.add_argument('output', type=Path, nargs='?')
    parser.add_argument('--draft', action='store_true')
    parser.add_argument('--docx', action='store_true')
    parser.add_argument('--font', default='Noto Sans CJK SC', help='Installed font family with Chinese glyphs for DOCX')
    args = parser.parse_args()
    try:
        if args.command == 'check':
            plan = json.loads(args.plan.read_text(encoding='utf-8'))
            report, _ = validate(plan, args.plan.resolve().parent)
            print(json.dumps(report, ensure_ascii=False, indent=2))
            return 0 if report['ready'] else 1
        if args.output is None:
            parser.error('build requires a new output directory')
        result = build(args.plan, args.output, args.draft, args.docx, args.font)
        print(json.dumps(result, ensure_ascii=False))
        return 0
    except (ValueError, OSError, ImportError, KeyError, TypeError) as exc:
        print(json.dumps({'ok': False, 'error': str(exc)}, ensure_ascii=False))
        return 1


if __name__ == '__main__':
    raise SystemExit(main())
