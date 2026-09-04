import copy
import importlib.util
import json
import shutil
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / 'scripts'))
import capture
import guide

HAS_PIL = importlib.util.find_spec('PIL') is not None
HAS_DOCX = importlib.util.find_spec('docx') is not None


class SamplingTests(unittest.TestCase):
    def test_sampling_covers_start_and_end_under_cap(self):
        times, adjusted = capture.sample_times(1000, max_frames=10)
        self.assertEqual(len(times), 10)
        self.assertEqual(times[0], 0)
        self.assertAlmostEqual(times[-1], 999.9)
        self.assertTrue(adjusted)

    def test_short_recording_and_nonzero_range(self):
        times, _ = capture.sample_times(.05)
        self.assertTrue(all(0 <= t < .05 for t in times))
        times, _ = capture.sample_times(20, start=10, end=15, interval=2)
        self.assertEqual(times, [10, 12, 14, 14.9])

    def test_bad_sampling_and_explicit_times(self):
        for kwargs in ({'interval': 0}, {'start': -1}, {'end': 100}, {'max_frames': 1},
                       {'times': [float('nan')]}, {'times': [10]}, {'times': [0, 1, 2], 'max_frames': 2}):
            with self.subTest(kwargs=kwargs), self.assertRaises(ValueError):
                capture.sample_times(10, **kwargs)
        self.assertEqual(capture.sample_times(10, times=[4, 2, 2])[0], [2, 4])

    def test_clock(self):
        self.assertEqual(guide.clock(3661.125), '01:01:01.125')

    def test_subprocess_no_shell_or_private_stderr(self):
        with self.assertRaises(ValueError) as result:
            capture.run([sys.executable, '-c', 'import sys; sys.stderr.write("secret"); sys.exit(1)'])
        self.assertNotIn('secret', str(result.exception))

    def test_help_works_from_other_directory(self):
        with tempfile.TemporaryDirectory() as tmp:
            for name in ('capture.py', 'guide.py'):
                result = subprocess.run([sys.executable, str(ROOT / 'scripts' / name), '--help'], cwd=tmp, capture_output=True)
                self.assertEqual(result.returncode, 0)
            self.assertFalse(list(Path(tmp).iterdir()))


@unittest.skipUnless(HAS_PIL, 'Pillow needed for image/export tests')
class GuideTests(unittest.TestCase):
    def setUp(self):
        from PIL import Image, PngImagePlugin
        self.temp = tempfile.TemporaryDirectory()
        self.base = Path(self.temp.name)
        index_dir = self.base / 'capture'
        index_dir.mkdir()
        pnginfo = PngImagePlugin.PngInfo()
        pnginfo.add_text('private', 'EXAMPLE SECRET')
        Image.new('RGB', (640, 360), 'white').save(index_dir / 'frame.png', pnginfo=pnginfo)
        self.index = {'schema': 'screen-frames/1', 'source': {'sha256': 'a'*64, 'duration': 10},
                      'frames': [{'id': 'f0001', 'at': 2, 'file': 'frame.png', 'width': 640, 'height': 360,
                                  'sha256': capture.digest(index_dir / 'frame.png')}]}
        self.write_index()
        self.plan = {'schema': 'screen-guide/1', 'title': '测试操作教程', 'purpose': '完成演示操作。',
                     'audience': '新人', 'scope': '演示页面', 'prerequisites': [], 'completion': ['出现完成提示。'],
                     'sources': {'main': 'capture/index.json'}, 'coverage_reviewed': True, 'questions': [],
                     'steps': [{'id': 'save', 'title': '保存', 'instruction': '点击保存按钮。', 'expected': '显示成功。',
                                'verified': True, 'privacy_reviewed': True,
                                'evidence': [{'source': 'main', 'frame': 'f0001', 'role': 'before',
                                              'redactions': [[10, 10, 100, 30]], 'boxes': [[300, 200, 100, 40]],
                                              'arrows': [[260, 170, 320, 220]]}]}]}

    def tearDown(self):
        self.temp.cleanup()

    def write_index(self):
        (self.base / 'capture/index.json').write_text(json.dumps(self.index), encoding='utf-8')

    def plan_path(self):
        path = self.base / 'guide.json'
        path.write_text(json.dumps(self.plan, ensure_ascii=False), encoding='utf-8')
        return path

    def test_valid_export_and_clean_public_directory(self):
        original = copy.deepcopy(self.plan)
        output = self.base / 'out'
        guide.build(self.plan_path(), output)
        self.assertTrue((output / 'COMPLETE.json').exists())
        text = (output / 'public/tutorial.html').read_text()
        self.assertIn('00:00:02.000', text)
        self.assertNotIn('核对单', text)
        self.assertNotIn('provenance', text)
        self.assertEqual(set(p.name for p in (output / 'public').iterdir()), {'tutorial.html', 'tutorial.md', 'images'})
        self.assertEqual(self.plan, original)

    def test_no_overwrite(self):
        output = self.base / 'out'
        guide.build(self.plan_path(), output)
        with self.assertRaises(FileExistsError):
            guide.build(self.plan_path(), output)

    def test_blockers_and_draft_separation(self):
        self.plan['questions'] = ['SECRET INTERNAL QUESTION']
        self.plan['steps'][0]['verified'] = False
        with self.assertRaises(ValueError):
            guide.build(self.plan_path(), self.base / 'out')
        self.assertFalse((self.base / 'out').exists())
        guide.build(self.plan_path(), self.base / 'draft', draft=True)
        self.assertIn('SECRET INTERNAL QUESTION', (self.base / 'draft/review.md').read_text())
        self.assertNotIn('SECRET INTERNAL QUESTION', (self.base / 'draft/public/tutorial.html').read_text())

    def test_privacy_and_coverage_gate(self):
        self.plan['coverage_reviewed'] = False
        self.plan['steps'][0]['privacy_reviewed'] = False
        report, _ = guide.validate(self.plan, self.base)
        self.assertTrue(report['valid'])
        self.assertFalse(report['ready'])
        self.assertEqual(len(report['blockers']), 2)

    def test_frame_hash_and_dimensions(self):
        self.index['frames'][0]['sha256'] = 'b'*64
        self.write_index()
        self.assertFalse(guide.validate(self.plan, self.base)[0]['valid'])
        self.index['frames'][0]['sha256'] = capture.digest(self.base / 'capture/frame.png')
        self.index['frames'][0]['width'] = 1
        self.write_index()
        self.assertFalse(guide.validate(self.plan, self.base)[0]['valid'])

    def test_frame_path_traversal(self):
        shutil.copyfile(self.base / 'capture/frame.png', self.base / 'private.png')
        self.index['frames'][0]['file'] = '../private.png'
        self.write_index()
        self.assertFalse(guide.validate(self.plan, self.base)[0]['valid'])

    def test_mixed_videos(self):
        index2 = copy.deepcopy(self.index)
        index2['source']['sha256'] = 'c'*64
        (self.base / 'capture/index2.json').write_text(json.dumps(index2))
        self.plan['sources']['second'] = 'capture/index2.json'
        self.assertFalse(guide.validate(self.plan, self.base)[0]['valid'])

    def test_shapes_and_crop(self):
        item = self.plan['steps'][0]['evidence'][0]
        for key, shape in [('boxes', [[-1, 0, 10, 10]]), ('redactions', [[0, 0, 0, 3]]),
                           ('arrows', [[0, 0, 800, 20]]), ('crop', [400, 100, 200, 200])]:
            changed = copy.deepcopy(self.plan)
            changed['steps'][0]['evidence'][0][key] = shape
            self.assertFalse(guide.validate(changed, self.base)[0]['valid'], key)
        item['crop'] = [0, 0, 500, 300]
        self.assertTrue(guide.validate(self.plan, self.base)[0]['valid'])

    def test_redaction_pixels_metadata_and_original_preserved(self):
        from PIL import Image
        original_hash = capture.digest(self.base / 'capture/frame.png')
        guide.build(self.plan_path(), self.base / 'out')
        with Image.open(self.base / 'out/public/images/step-001-01.png') as image:
            self.assertEqual(image.getpixel((30, 20)), (17, 17, 17))
            self.assertNotIn('private', image.info)
            self.assertEqual(image.getpixel((300, 220)), (220, 38, 38))
        self.assertEqual(capture.digest(self.base / 'capture/frame.png'), original_hash)

    def test_escape_html_and_markdown(self):
        self.plan['steps'][0]['instruction'] = '<script>alert(1)</script> [click](https://example.test)'
        guide.build(self.plan_path(), self.base / 'out')
        rendered = (self.base / 'out/public/tutorial.html').read_text()
        self.assertNotIn('<script>', rendered)
        self.assertIn('&lt;script&gt;', rendered)
        self.assertIn('\\[click\\]', (self.base / 'out/public/tutorial.md').read_text())

    def test_bad_inputs(self):
        for plan in (None, [], {}, {'schema': 'screen-guide/1', 'sources': [], 'steps': [None]}):
            self.assertFalse(guide.validate(plan, self.base)[0]['valid'])
        self.plan['steps'][0]['evidence'] = [{'source': [], 'frame': []}]
        self.assertFalse(guide.validate(self.plan, self.base)[0]['valid'])

    def test_missing_duplicate_and_reverse_steps(self):
        self.plan['steps'][0]['evidence'][0]['frame'] = 'missing'
        self.assertFalse(guide.validate(self.plan, self.base)[0]['valid'])
        self.plan['steps'][0]['evidence'][0]['frame'] = 'f0001'
        self.plan['steps'].append(copy.deepcopy(self.plan['steps'][0]))
        self.assertFalse(guide.validate(self.plan, self.base)[0]['valid'])
        self.plan['steps'][1]['id'] = 'earlier'
        earlier = copy.deepcopy(self.index['frames'][0])
        earlier.update(id='earlier', at=1)
        self.index['frames'].append(earlier)
        self.write_index()
        self.plan['steps'][1]['evidence'][0]['frame'] = 'earlier'
        self.assertFalse(guide.validate(self.plan, self.base)[0]['ready'])

    @unittest.skipUnless(HAS_DOCX, 'python-docx needed')
    def test_docx_text_images_and_privacy(self):
        from docx import Document
        guide.build(self.plan_path(), self.base / 'out', docx=True)
        doc = Document(self.base / 'out/public/tutorial.docx')
        self.assertEqual(len(doc.inline_shapes), 1)
        self.assertTrue(any('点击保存按钮' in p.text for p in doc.paragraphs))
        self.assertFalse(any('内部核对' in p.text for p in doc.paragraphs))
        self.assertFalse(doc.styles.element.xpath('.//w:pBdr'))


@unittest.skipUnless(HAS_PIL and shutil.which('ffmpeg') and shutil.which('ffprobe'), 'FFmpeg/Pillow needed')
class CaptureIntegrationTests(unittest.TestCase):
    def test_real_decode_audio_absence_and_capped_range(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            video = root / 'test clip 中文.mp4'
            capture.run([capture.executable('ffmpeg'), '-v', 'error', '-nostdin', '-n', '-f', 'lavfi',
                         '-i', 'color=c=blue:s=320x180:r=10:d=2', '-c:v', 'libx264', str(video)])
            result = capture.capture(video, root / 'frames', interval=.1, max_frames=3, audio=True)
            self.assertTrue(result['cap_adjusted'])
            self.assertFalse(result['audio_extracted'])
            index = json.loads((root / 'frames/index.json').read_text())
            self.assertEqual(len(index['frames']), 3)
            self.assertEqual(index['frames'][0]['width'], 320)
            self.assertGreater(index['frames'][-1]['at'], 1.8)
            with self.assertRaises(FileExistsError):
                capture.capture(video, root / 'frames', times=[0])

    def test_audio_extract_with_offset(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            video = root / 'audio.mp4'
            capture.run([capture.executable('ffmpeg'), '-v', 'error', '-nostdin', '-n', '-f', 'lavfi',
                         '-i', 'color=c=blue:s=320x180:r=10:d=2', '-f', 'lavfi',
                         '-i', 'sine=frequency=440:duration=2', '-c:v', 'libx264', '-c:a', 'aac', str(video)])
            result = capture.capture(video, root / 'frames', start=.5, end=1.5, times=[.6], audio=True)
            self.assertTrue(result['audio_extracted'])
            import wave
            with wave.open(str(root / 'frames/audio.wav')) as audio:
                self.assertEqual(audio.getframerate(), 16000)
                self.assertEqual(audio.getnchannels(), 1)
                self.assertAlmostEqual(audio.getnframes() / audio.getframerate(), 1, places=1)


if __name__ == '__main__':
    unittest.main()
